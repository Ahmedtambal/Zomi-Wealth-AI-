from turtle import st
import openai
from docx import Document
from datetime import datetime
import pdfplumber
import mimetypes
import PyPDF2
import json
import os
import re  # Ensure this is included
import openai
import pytesseract
from PIL import Image
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


# Get the API key from environment variable
openai.api_key=""
UPLOAD_FOLDER = "uploaded_docs"  # Ensure it's defined globally


# Specify the full path to Tesseract executable if not in PATH
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def clean_json_response(response_content: str) -> str:
    """Clean GPT response containing JSON data"""
    # Remove markdown code blocks
    cleaned = re.sub(r'^```json\s*|\s*```$', '', response_content, flags=re.DOTALL)
    # Remove non-UTF8 characters and special spaces
    cleaned = cleaned.encode('utf-8', 'ignore').decode('utf-8')
    # Remove left-to-right marks and other special characters
    cleaned = re.sub(r'[\u200e\u200f]', '', cleaned)
    return cleaned.strip()

def parse_json_response(response_content: str, context: str = "") -> dict:
    """Parse and validate JSON response with error handling"""
    try:
        return json.loads(response_content)
    except json.JSONDecodeError as e:
        error_msg = (
            f"JSON parsing failed in {context}:\n"
            f"Error: {str(e)}\n"
            f"Original content: {repr(response_content)}\n"
        )
        raise ValueError(error_msg) from e
    
def extract_text_from_file(self, file_path):
        """
        Extract text from a file (PDF or Image).
        """
        _, file_extension = os.path.splitext(file_path)
        if file_extension.lower() == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_extension.lower() in [".png", ".jpg", ".jpeg"]:
            return self.extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

def extract_text_from_image(image_path):
    """
    Extract text from an image using Tesseract OCR.
    
    Args:
    - image_path (str): Path to the image file.

    Returns:
    - str: Extracted text from the image.
    """
    # Load the image
    image = Image.open(image_path)

    # Extract text using Tesseract OCR
    extracted_text = pytesseract.image_to_string(image)

    return extracted_text


def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF file using pdfplumber.
    """
    try:
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages])
        return text
    except Exception as e:
        raise ValueError("Error extracting text from PDF: " + repr(e))

def extract_text_from_multiple_files(file_paths):
    """
    Extract text from multiple files and return them as a dictionary.
    Args:
    - file_paths (list of str): List of file paths.

    Returns:
    - dict: A dictionary where keys are file names and values are extracted text.
    """
    extracted_data = {}
    for file_path in file_paths:
        file_name = os.path.basename(file_path)
        extracted_data[file_name] = extract_text_from_file(file_path)
    return extracted_data    


def extract_risk_details(file_path):
    """
    Extract risk level, type, first sentence, and last sentence from an uploaded image or document using OCR.
    
    Args:
    - file_path: Path to the uploaded image or document.
    
    Returns:
    - A dictionary with risk details (level, type, first sentence, last sentence).
    """
    try:
        # Use Tesseract OCR to extract text from the uploaded image
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)

        # Extract relevant details
        level_of_risk = re.search(r"Risk Level\s(\d+)", text)
        risk_level = re.search(r"Risk Type:\s([\w\s-]+)", text)

        # Extract first and last sentence from the "Definition of [Risk Type]" section
        definition_section = re.search(r"Definition of [\w\s-]+:(.+)", text, re.DOTALL)
        if definition_section:
            sentences = definition_section.group(1).strip().split(".")
            first_sentence = sentences[0].strip() + "." if len(sentences) > 0 else ""
            last_sentence = sentences[-1].strip() + "." if len(sentences) > 1 else ""
        else:
            first_sentence = ""
            last_sentence = ""

        return {
            "level_of_risk": level_of_risk.group(1) if level_of_risk else "Unknown",
            "risk_level": risk_level.group(1).strip() if risk_level else "Unknown",
            "first_sentence": first_sentence,
            "last_sentence": last_sentence,
        }
    except Exception as e:
        raise ValueError("Error extracting risk details: " + repr(e))
    
def process_plan_report(uploaded_file):
    """
    Process the uploaded plan report, extract text, and generate table data.
    """
    # Save the uploaded file
    file_path = save_uploaded_file(uploaded_file, UPLOAD_FOLDER)

    # Step 1: Extract text from the file
    extracted_text = extract_text_from_pdf(file_path)  # Assuming PDF for now

    # Step 2: Use GPT to process the extracted text and get dynamic data
    gpt_response = extract_plan_details_with_gpt(extracted_text)

    try:
        # Parse the JSON response
        table_data = json.loads(gpt_response)
        if not isinstance(table_data, list):
            raise ValueError("Parsed GPT response is not a list of dictionaries.")
    except json.JSONDecodeError as e:
        raise ValueError("Error decoding GPT response: " + repr(e) + "\nResponse: " + gpt_response)

    return table_data




def extract_client_details_with_gpt(factfinding_text):
    prompt = f"""
    You are an AI assistant tasked with extracting specific client details from a FactFinding report.

    **Objective**:
    Analyze the provided FactFinding report and extract the following details to populate placeholders in a financial document:

    **Placeholders**:
    - Full name: Combine Title and Surname
    - Address: Full multiline address with postal code
    - Today's date: Current date in "9th January 2025" format
    - Salutation: "Dear [Forename]," format

    **FactFinding Report**:
    {factfinding_text}

    **Expected JSON Format**:
    {{
      "Full name": "",
      "Address": "",
      "Today’s date": "",
      "salutation": ""
    }}
    """
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        raw_content = response["choices"][0]["message"]["content"]
        cleaned_content = clean_json_response(raw_content)
        return parse_json_response(cleaned_content, "client details extraction")
    except Exception as e:
        error_msg = f"Client details error: {str(e)}"
        if 'raw_content' in locals():
            error_msg += f"\nRaw response: {raw_content}"
        raise ValueError(error_msg)
    

def generate_current_situation(factfinding_text):
    """
    Use OpenAI to generate a detailed 'Current Situation' section based on the FactFinding report.
    """
    prompt = f"""
    You are an assistant tasked with creating a "Current Situation" summary for a financial report.

    **Objective**:
    1. Analyze the provided FactFinding report and extract all relevant details.
    2. Write a detailed "Current Situation" section summarizing the client's financial and personal details.
    3. Write the "Current Situation" section as bullet points only. Do not include any headings, introductions.

    **Instructions**:
    - Use UK grammar, language and date format, also dont use 'z' in words use 's' where applicable for example (e.g., "realise" instead of "realize").
   - 1. Use British English spelling conventions:
        - Words ending in "ise" (e.g., "prioritise" instead of "prioritize").
        - Words like "colour", "favour", and "neighbour" (instead of "color", "favor", and "neighbor").
        - Words like "centre" and "metre" (instead of "center" and "meter").
        - Words like "travelling" and "cancelled" with double "l" (instead of "traveling" and "canceled").
   - 2.Use British grammar and punctuation:
        - Use single quotation marks (' ') for quotes instead of double quotation marks (" ").
        - Place punctuation outside quotation marks unless it's part of the quoted text.
        - Refer to collective nouns (e.g., "team", "government") as plural where appropriate (e.g., "The team are ready").
   - 3.Use British terminology:
        - Use "flat" instead of "apartment", "lift" instead of "elevator", and "petrol" instead of "gasoline".
        - Use "holiday" instead of "vacation" and "autumn" instead of "fall".    
    - Write the section in professional bullet points, but keep it conversational by using "You" at the beginning of some sentences.
    - Make sure to mention the part about include monthly growth and monthly expenditure and remain in one line for example (• You have a monthly gross income of £2,700.00 and a monthly expenditure of £1,670.00, leaving you with a monthly surplus of £1,030.00.).
    - Include point about dependants and Wife details and their financial details if they are found.
    - When extracting dependents' details, exclude the name of the person who took the notes (e.g., "16 Sep 2024 - Alex Armstrong"). Focus only on the actual dependents' names and their details.
    - Whenever you mention the phrase "You should always have 3-6 months worth of expenditure," always calculate and include the range for 3-6 months of monthly expenditure, formatted as (£<3 months value>-£<6 months value>). Use the extracted monthly expenditure for this calculation(3 multiple by monthly expenditure - 6 multiple by monthly expenditure ).
    - Include line about Protection detalis (e.g,  Home insurance, car insurance,..... .)
    - Include sufficient length and detail, following this example:

    **Example "Current Situation" Section**:
    
    Current Situation

    When we last met, you made me aware of your current situation:

    • Chris, you are 68 years old, co-habiting with your partner and in good health. 
    • You retired in January 2020 having previously worked in IT.  
    • You own your house outright which is worth approximately £555,000.00.
    • You are in receipt of your full state pension (£930.00 gross per month), Zurich Financial Services final salary pension (£640.00 gross per month) and your NHS pension (£890.00 gross per month).
    • You withdraw £800.00 gross per month from your tax-free cash entitlement from your Royal London Personal Pension. This will be exhausted soon. 
    • You have no debts or other liabilities.
    • You have a monthly gross income of £3,260.00 and a monthly expenditure of £2,410.00, leaving you with a monthly disposable of £850.00. 
    • You have no financial dependents. 
    • You have £39,000.00 in cash reserves. This is a sufficient emergency fund. You should always have 3-6 months worth of expenditure in an easy access bank account for emergencies (£7,230.00-£14,460.00). You may want to invest any over this amount to get better returns than cash in the bank. 
    • You have drafted a Will and Power of Attorney, and they are both up to date.

    **Your Task**:
    1. Ensure your output matches the tone, structure, and level of detail of the example provided above.
    2. Extract details from the FactFinding report, including:
       - Personal details (age, marital status, and health).
       - Dependants and Wife details and their financial details if they are found.
       - Include point about dependants and Wife details and their financial details if they are found.
       - Retirement details and previous occupation , also for his wife it was found in one line with his detalis.
       - Property ownership and approximate value.
       - Pensions, incomes, and cash withdrawals.
       - Include point about monthly growth and monthly expenditure in one line, you can use the Incomes and Expenses table (The Sum of each), use the word growth and exenditure when talking about it.
       - Emergency funds and recommendations for improvement.
       - Succession planning (Will and Power of Attorney).
       - Protection
       - Any other relevant details.

    **FactFinding Report**:
    {factfinding_text}

    **Output**:
    Write a "Current Situation" section that:
    1. Matches the tone and format of the example.
    2. Includes sufficient detail to make the section comprehensive and professional.
    3. Avoids mentioning specific details about thier invstment knowledge and experience.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    return response["choices"][0]["message"]["content"]


def generate_priorities_and_objectives(factfinding_text):
    """
    Use OpenAI to generate a detailed 'Priorities and Objectives' section based on the FactFinding report.
    """
    prompt = f"""
    You are an assistant tasked with creating a "Priorities and Objectives" section for a financial report.

    **Objective**:
    1. Analyze the provided FactFinding report and extract relevant details fro m the Objectives table.
    2. Personalize the "Priorities and Objectives" section by integrating other relevant information in the report, such as financial circumstances, retirement plans, health, and family situation.
    3. Write the "Priorities and Objectives" section as a bullet points or numbers for example (e.g., 1. lab akbdh sj h , 2. hushfjfhsuhfivn;k , 3.h osijck,........). Do not include headings or numbered lists.

    **Instructions**:
    - Use UK grammar, language and date format, also dont use 'z' in words use 's' where applicable for example (e.g., "realise" instead of "realize").
    - 1. Use British English spelling conventions:
        - Words ending in "ise" (e.g., "prioritise" instead of "prioritize").
        - Words like "colour", "favour", and "neighbour" (instead of "color", "favor", and "neighbor").
        - Words like "centre" and "metre" (instead of "center" and "meter").
        - Words like "travelling" and "cancelled" with double "l" (instead of "traveling" and "canceled").
    - 2.Use British grammar and punctuation:
        - Use single quotation marks (' ') for quotes instead of double quotation marks (" ").
        - Place punctuation outside quotation marks unless it's part of the quoted text.
        - Refer to collective nouns (e.g., "team", "government") as plural where appropriate (e.g., "The team are ready").
    - 3.Use British terminology:
        - Use "flat" instead of "apartment", "lift" instead of "elevator", and "petrol" instead of "gasoline".
        - Use "holiday" instead of "vacation" and "autumn" instead of "fall".   
    - Focus on the client's primary and secondary financial objectives, integrating personal context to make it highly personalized.
    - Use a professional tone and structure, with clear and specific details.
    - Avoid unnecessary repetition or vague language.
    - Ensure that the content remains cohesive and logical.
    - Write the "Priorities and Objectives" section as a bullet points or numbers for example (e.g., 1. lab akbdh sj h , 2. hushfjfhsuhfivn;k , 3.h osijck,........).


    **Your Task**:
    1. Ensure your output matches the tone, structure, and level of detail of the example provided above.
    2. Extract details from the FactFinding report, including:
       - Financial objectives (e.g., maintaining living standards, creating income, building capital).
       - Any specific income goals or plans (e.g., covering monthly expenses or specific purchases).
       - Retirement plans or strategies (e.g., utilizing pensions, maintaining capital for long-term care).
       - Any family or personal considerations that impact objectives (e.g., no descendants, focus on enjoying funds in the client's lifetime).
       - Relevant details from the Objectives table, such as priority goals, timeframes, and preferences.
    3. Write the section in a professional and concise manner.

    **FactFinding Report**:
    {factfinding_text}

    **Output**:
    Write a "Priorities and Objectives" section that:
    1. Matches the tone and structure of the example provided above.
    2. Includes sufficient personalization based on the Objectives table and other relevant details.
    3. Avoids unnecessary repetition or vague statements.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    return response["choices"][0]["message"]["content"]


def extract_details_with_gpt(extracted_text):
    prompt = f"""
    Extract these details from the text:
    1. First Name
    2. Risk Level (1-5)
    3. Risk Type (e.g., Cautious)
    4. First sentence of risk definition
    5. Last sentence of risk definition

    **Text**:
    {extracted_text}

    **Expected JSON Format**:
    {{
        "first_name": "",
        "risk_level": "",
        "risk_type": "",
        "first_sentence": "",
        "last_sentence": ""
    }}
    """
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        raw_content = response["choices"][0]["message"]["content"]
        cleaned_content = clean_json_response(raw_content)
        return parse_json_response(cleaned_content, "risk details extraction")
    except Exception as e:
        error_msg = f"Risk details error: {str(e)}"
        if 'raw_content' in locals():
            error_msg += f"\nRaw response: {raw_content}"
        raise ValueError(error_msg)
      

def extract_plan_details_with_gpt(extracted_text):
    prompt = f"""
    Extract plan details from this financial report:

    **Required Fields**:
    - Provider name
    - Plan Number
    - Plan Type
    - Current Value

    **Example Output**:
    [
        {{
            "Provider": "Royal London",
            "Plan Number": "3226624",
            "Plan Type": "Personal Pension",
            "Current Value": "£211,058.25"
        }}
    ]

    **Report Text**:
    {extracted_text}
    """
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        raw_content = response["choices"][0]["message"]["content"]
        cleaned_content = clean_json_response(raw_content)
        return parse_json_response(cleaned_content, "plan details extraction")
    except Exception as e:
        error_msg = f"Plan details error: {str(e)}"
        if 'raw_content' in locals():
            error_msg += f"\nRaw response: {raw_content}"
        raise ValueError(error_msg)
       
    
def generate_pension_review_section(product_report_text):
    """
    Use OpenAI to generate a detailed 'Review of Existing Royal London Personal Pension' section
    based on the product report.
    """
    prompt = f"""
    You are a financial assistant tasked with creating a "Review of Existing Royal London Personal Pension" section for a financial report.

    **Objective**:
    1. Analyze the provided product report and extract all relevant details.
    2. Write a detailed and personalized "Review of Existing Royal London Personal Pension" section based on the client's financial details.
    3. Structure the section in a conversational tone, starting sentences with "You" and addressing the reader directly.

    **Instructions**:
    - Use UK grammar, language and date format, also dont use 'z' in words use 's' where applicable for example (e.g., "realise" instead of "realize").
    - Include the following details:
      - Regular and one-off withdrawals, distinguishing between taxable and tax-free.
      - Historical and recent contributions, including any profit share awards.
      - The split between the income release account and savings account.
      - Availability of tax-free cash and its current status.
      - Annual management charges, transaction costs, and rebates.
      - It's important to not mention any detaleid information about the invesment protfolio, just the names and the objectives, keep it short.
    - Write clearly and concisely in full paragraphs, avoiding bullet points.
    - Use a conversational tone that feels direct and engaging, starting with "You" where appropriate.
    - 1. Use British English spelling conventions:
        - Words ending in "ise" (e.g., "prioritise" instead of "prioritize").
        - Words like "colour", "favour", and "neighbour" (instead of "color", "favor", and "neighbor").
        - Words like "centre" and "metre" (instead of "center" and "meter").
        - Words like "travelling" and "cancelled" with double "l" (instead of "traveling" and "canceled").
    - 2.Use British grammar and punctuation:
        - Use single quotation marks (' ') for quotes instead of double quotation marks (" ").
        - Place punctuation outside quotation marks unless it's part of the quoted text.
        - Refer to collective nouns (e.g., "team", "government") as plural where appropriate (e.g., "The team are ready").
    - 3.Use British terminology:
        - Use "flat" instead of "apartment", "lift" instead of "elevator", and "petrol" instead of "gasoline".
        - Use "holiday" instead of "vacation" and "autumn" instead of "fall".

Ensure all responses adhere to these conventions.
    **Example Section**:
    You are currently withdrawing a total of £1,200.00 per month from your pension, amounting to £14,400.00 annually. This entire amount is taxable, as you have already utilized your tax-free cash entitlement. Additionally, you made a one-off taxable withdrawal of £10,750.00 during this review period, bringing your total withdrawals to £25,150.00. Since the inception of your plan, you have withdrawn £131,953.95, with £45,983.94 being tax-free and £85,970.01 taxable.

    Your pension plan is currently valued at £124,029.82, entirely held within the Income Release Account, with no funds in the Savings Account. Consequently, there is no tax-free cash available at this time. It is important to consider the implications of continuing taxable withdrawals, as this will affect your overall tax liability and could potentially trigger the Money Purchase Annual Allowance (MPAA), limiting future pension contributions to £10,000.00 per annum.

    The ongoing Annual Management Charge for your plan is 0.45% per annum, equating to £594.12 for this review period. This charge is indicative, as different rates apply to various contributions. Additionally, your investment return accounts for transaction costs of £39.48. You have also benefited from a ProfitShare award of £209.64, which contributes to reducing your overall charges.

    Your pension is invested in the Royal London Governed Portfolio Growth, designed to achieve above-inflation growth with a medium to high-risk profile. The portfolio is diversified across 12 funds, with significant allocations in RLP Global Managed (63.23%) and RLP Property (10.21%), among others. During this review period, 19 changes were made to your portfolio to optimize asset allocation in response to market conditions. The annualized performance of your plan over this period was 10.2%, and since inception, it has been 6.0%.

    **Your Task**:
    1. Extract details from the product report to generate a section similar to the example above.
    2. Ensure the tone and structure match the example provided.

    **Product Report**:
    {product_report_text}

    **Output**:
    Write a "Review of Existing Royal London Personal Pension" section that:
    - Matches the conversational tone and format of the example.
    - Includes sufficient detail to make the section comprehensive and professional.
    """

    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    return response["choices"][0]["message"]["content"]    
    
def generate_safe_withdrawal_rate_section(plan_report_text):
    """
    Use OpenAI to generate a detailed 'Safe Withdrawal Rate (SWR)' section
    based on the plan report.

    Args:
        plan_report_text (str): Extracted text from the plan report.

    Returns:
        str: The generated SWR section or an empty string if no withdrawals are detected.
    """
    # Ensure your OpenAI API key is set in the environment variables
    
    prompt = f"""
    You are a financial assistant tasked with creating a "Safe Withdrawal Rate (SWR)" section for a financial report.

    **Objective**:
    1. Analyze the provided plan report text to determine if the client is withdrawing money from their investments (e.g., pensions, savings, other investments).
    2. If withdrawals are present, extract the following details:
       - Monthly withdrawal amount.
       - Annual withdrawal rate.
       - Total portfolio value.
    3. Generate a detailed and personalized "Safe Withdrawal Rate (SWR)" section based on the extracted information.
    4. Use a conversational tone that feels direct and engaging, starting sentences with "You" where appropriate.

    **Instructions**:
    - Use UK grammar, language and date format, also dont use 'z' in words use 's' where applicable for example (e.g., "realise" instead of "realize").
    - 1. Use British English spelling conventions:
        - Words ending in "ise" (e.g., "prioritise" instead of "prioritize").
        - Words like "colour", "favour", and "neighbour" (instead of "color", "favor", and "neighbor").
        - Words like "centre" and "metre" (instead of "center" and "meter").
        - Words like "travelling" and "cancelled" with double "l" (instead of "traveling" and "canceled").
    - 2.Use British grammar and punctuation:
        - Use single quotation marks (' ') for quotes instead of double quotation marks (" ").
        - Place punctuation outside quotation marks unless it's part of the quoted text.
        - Refer to collective nouns (e.g., "team", "government") as plural where appropriate (e.g., "The team are ready").
    - 3.Use British terminology:
        - Use "flat" instead of "apartment", "lift" instead of "elevator", and "petrol" instead of "gasoline".
        - Use "holiday" instead of "vacation" and "autumn" instead of "fall".    
    - If withdrawals are detected:
      - Calculate the current withdrawal rate as (annual withdrawals / total portfolio value) * 100, dont write the equation just the results.
      - Compare the current withdrawal rate to the generally accepted safe withdrawal rate of 4.00%.
      - Highlight the risks associated with withdrawing more than the recommended rate.
    - If no withdrawals are detected, respond with "No withdrawals detected."

    - Write clearly and concisely in full paragraphs, avoiding bullet points.

    **Example Section**:
    Safe Withdrawal Rate (SWR)

    You're taking a taxable income of £1,200.00 from your plan every month.

    The generally accepted safe withdrawal rate for retirement income is approximately 4.00% per year of the portfolio’s value.

    Your current withdrawal rate is: 11.61%.

    Withdrawing more than 4.00% annually may deplete your investment faster than anticipated, especially during periods of market volatility.

    With a withdrawal rate exceeding 4.00%, there is an increased risk that your investment may not last throughout your retirement. This could result in a shortfall in later years, reducing your ability to meet essential expenses. We recommend regular reviews to ensure your withdrawals remain sustainable.

    **Plan Report Text**:
    {plan_report_text}

    **Output**:
    Write a "Safe Withdrawal Rate (SWR)" section that:
    - Follows the structure and tone of the example provided.
    - Includes all relevant extracted details.
    - Provides a clear comparison between the client's withdrawal rate and the recommended rate.
    """

    try:
        # Make a call to the OpenAI API
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a financial advisor assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,  # Lower temperature for more deterministic output
            max_tokens=600,    # Adjust as needed to capture detailed responses
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )

        # Extract the response text
        generated_text = response.choices[0].message['content'].strip()

        if generated_text == "No withdrawals detected.":
            return ""  # No SWR section needed
        else:
            return generated_text

    except Exception as e:
        print("Error generating SWR section: " + repr(e))
        return ""
    
    
def extract_fund_performance_with_gpt(extracted_text):
    """
    Simplified version that makes GPT handle all calculations
    """
    prompt = f"""
    Analyze this fund performance data and return JSON with:
    1. Yearly performance percentages
    2. Benchmark comparisons
    3. Cumulative 5-year sum (calculated as simple sum of yearly percentages)
    
    Rules:
    - Use EXACTLY this format:
    [
        {{
            "Fund": "Fund Name",
            "Year 1": "X%",
            "Year 2": "X%",
            "Year 3": "X%", 
            "Year 4": "X%",
            "Year 5": "X%",
            "Cumulative (5 YR)": "X%",
            "Benchmark": {{
                "Year 1": "X%",
                "Year 2": "X%",
                "Year 3": "X%",
                "Year 4": "X%",
                "Year 5": "X%",
                "Cumulative (5 YR)": "X%"
            }}
        }}
    ]
    - Only return raw JSON without any additional text/comments
    - Use "N/A" for missing data
    - Never include Markdown formatting (no ```json)
    
    Text to analyze:
    {extracted_text}
    """

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        
        # Get and clean response
        raw_content = response["choices"][0]["message"]["content"]
        cleaned_content = raw_content.strip().replace('```json', '').replace('```', '')
        
        return json.loads(cleaned_content)
        
    except json.JSONDecodeError as e:
        error_msg = f"Failed to parse GPT response: {str(e)}\nRaw response: {raw_content}"
        raise ValueError(error_msg)
    except Exception as e:
        raise ValueError(f"Fund performance error: {str(e)}")

def extract_dark_star_performance_with_gpt(extracted_text):
    """
    Extract Dark Star fund performance data from the extracted text using GPT.
    """
    prompt = f"""
    You are an AI assistant tasked with extracting fund performance details from a financial report.
    Analyze the text below and return a JSON response similar to this format:
    [
        {{
            "Fund": "Dark Star Asset Management Balanced Plus",
            "Year 1": "15%",
            "Year 2": "9.7%",
            "Year 3": "6.5%",
            "Year 4": "4.3%",
            "Year 5": "N/A",
            "Benchmark": {{}},
            "Cumulative (5 YR)": "33.9%"
        }}
    ]

    If no benchmark is provided in the text, leave the "Benchmark" field empty.
    **Important Instructions**:
        - Use the "YTD" (Year-to-Date) data for yearly performance.
        - Map the most recent year as Year 1, the next as Year 2, and so on.
        - If a year has no data, use "N/A".
        - Cumulative performance should only include available years and should be calculated as the sum of those percentages.

    **Example**:
        If the text contains:
        ```
        % Jan Feb Mar Apr May Jun Jul Aug Sep Oct Nov Dec YTD
        2024 -1.0 1.8 2.6 -1.0 0.3 1.4 0.3 0.0 1.4 9.1
        2023 - - - - - - 1.2 -1.5 -0.5 -1.8 3.5 3.8 4.7
        ```
        The response should be:
        [
            {{
                "Fund": "Dark Star Asset Management Balanced Plus",
                "Year 1": "9.1%",
                "Year 2": "4.7%",
                "Year 3": "N/A",
                "Year 4": "N/A",
                "Year 5": "N/A",
                "Cumulative (5 YR)": "11.8%"
            }}
        ]


    Text:
    {extracted_text}
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        raw_content = response["choices"][0]["message"]["content"]
        
        # Add cleaning
        cleaned_content = re.sub(r'^```json\s*|\s*```$', '', raw_content, flags=re.DOTALL)
        
        return json.loads(cleaned_content)
    except json.JSONDecodeError as e:
        raise ValueError(f"Dark Star JSON error: {e}\nResponse: {raw_content}")
    

def extract_sap_comparison_with_gpt(extracted_text):
    """
    Extract comparison details and age from the SAP report using GPT.
    """
    prompt = f"""
    You are an AI assistant tasked with extracting a comparison table and the relevant age mentioned in a financial SAP report.

    Analyze the following text and return a JSON response like this:
    {{
        "Age": 80,
        "Table": {{
            "Assumed Growth Rates": ["0%", "2.94%", "5.88%"],
            "Existing Schemes": ["£118,972.00", "£155,558.00", "£201,866.00"],
            "P1 Pension Account": ["£111,000.00", "£145,000.00", "£189,000.00"],
            "Rate of Return Required from P1": ["+0.69%", "+0.72%", "+0.74%"],
            "Effect on Fund if Moved to P1": ["-6.19%", "-6.19%", "-6.19%"],
            "Reduction in Yield if Moved to P1": ["1.20%", "1.20%", "1.20%"]
        }}
    }}

    Text:
    {extracted_text}

    Ensure that the extracted age matches the age mentioned in the heading (e.g., "Comparison at Age 80") and that the table is correctly formatted.
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        raw_content = response["choices"][0]["message"]["content"]
        
        # Add cleaning
        cleaned_content = re.sub(r'^```json\s*|\s*```$', '', raw_content, flags=re.DOTALL)
        
        return json.loads(cleaned_content)
    except json.JSONDecodeError as e:
        raise ValueError(f"SAP comparison JSON error: {e}\nResponse: {raw_content}")

def extract_annuity_quotes_with_gpt(extracted_text):
    """
    Use GPT to extract annuity quotes details from the extracted text and return a plain text response.
    """
    prompt = f"""
        ### Instructions:
        Extract the details of annuity quotes from the provided text and format the output as follows:

        Example Output:
        Quote 1:
        - Purchase Amount: £124,030
        - Monthly Amount: £854
        - Yearly Amount: £10,250
        - Yearly Increase: None

        Quote 2:
        - Purchase Amount: £124,030
        - Monthly Amount: £603
        - Yearly Amount: £7,242
        - Yearly Increase: Retail Price Index (RPI)

        Quote 3:
        - Purchase Amount: £124,030
        - Monthly Amount: £668
        - Yearly Amount: £8,018
        - Yearly Increase: 3.00%

        ### Notes:
        1. **Purchase Amount**:
        - Locate the number next to the text `pension pot`.
        - For example, if you see `£124,030 pension pot`, extract `£124,030`.

        2. **Monthly Amount**:
        - Find the number followed by the word `monthly` under the **Your Income** section.
        - For example, if you see `£854 monthly`, extract `£854`.

        3. **Yearly Amount**:
        - Find the number followed by the word `yearly` under the **Your Income** section.
        - For example, if you see `£10,250 yearly`, extract `£10,250`.

        4. **Yearly Increase**:
        - Look for the yearly increase information in the **Your Choices** section.
        - Extract as:
            - `"None"` for `No annual increase`.
            - `"Retail Price Index (RPI)"` for `Increase by RPI`.
            - `"3.00%"` for `Increase 3% per year`.

        ---

        Text to Analyze:
    {extracted_text}
    """
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
    )
    # Return the formatted plain text output
    return response["choices"][0]["message"]["content"]


def extract_fund_comparison_with_gpt(fund1_text, fund2_file1_text, fund2_file2_text):
    """
    Generate a GPT response for fund comparison in the specified short format.
    """
    prompt = f"""
    ### Role
    You are a financial data extraction system. Analyze the provided text for Royal London and P1 funds, extract specific percentages and values, and populate the template below. 

    ### Input Text
    **Royal London Details**:
    {fund1_text}

    **P1 Details**:
    {fund2_file1_text}
    {fund2_file2_text}

    ### Tasks
    1. **Extract Values**:
       - **Plan value**: Find the value matching the end date of the review period (look for "Review dates:").
       - **Royal London Charges**:
         - Weighted Fund Charge: Extract from phrases like "equivalent to X% of the value of your plan each year."
         - Platform Charge: Search for "Platform" or "[Provider Name] charges". Default to 0.0% if missing.
         - Ongoing Advice Fee: Default to 0.50% if missing.always 0.50%.
         - Discretionary Fund Manager Charge: 0.0% if not mentioned (look for "Discretionary").
         - Drawdown Fee: Search for "Drawdown" or "Product".
         - ProfitShare: Always -0.15%.
       - **P1 Charges**:
         - check the text after the section " Itemisation of Service Charges  Charges paid for all services in the first year:"
         - Discretionary Fund Manager Charge: Refreed to as "Discretionary Manager" fees so look for that. 
         - Extract the same categories. Use defaults where specified (e.g., Weighted Fund Charge = 0.44% if missing).
         - Drawdown Fee: Search for "Drawdown" or "Product".
         - ProfitShare: Extract from the text, if is not found 0.0 . IS NOT THE SAME AS ROYAL LONDON.  
         - Platform Charge: Search for "Platform" or "[Provider Name] charges".

    2. **Calculate Totals**:
       - Multiply all percentages by Royal London’s Plan Value to get £ amounts.
       - Sum percentages and £ amounts for **Total Annual Ongoing Charges**.

    3. **Compare Funds**:
       - Compute differences (P1 - Royal London) for % and £.
       - Generate a dynamic comparison statement.

    ### Output Format
    Return **ONLY** the populated template below. No explanations, placeholders, or markdown:

    **Step 2: Calculate £ Amounts**
- For **both funds**, multiply each percentage by the **Royal London Plan Value** (e.g., 0.45% * £124,029.82 = £558.13).
- Sum all percentages and £ amounts for **Total Annual Ongoing Charges**. then save the values becuz you will need them LATER TO MAP (£[Value]) for each value(weighted fun charge, platform charge, ....).

**Step 3: Compare Funds**
- Compute differences: P1 Total %/£ - Royal London Total %/£.
- Generate a comparison statement (e.g., "P1 is 0.10% cheaper...").

**Step 4: Output Format**
Return **ONLY** this template with values filled:

---
Plan value = £[Value]  
Weighted Fund Charge % = [Value]% (£[Value])  
Platform Charge % = [Value]% (£[Value])  
Ongoing Advice Fee % = 0.50% (£[Value])  
Discretionary Fund Manager Charge % = [Value]% (£[Value])  
Drawdown Fee % = [Value]% (£[Value])  
ProfitShare % = -0.15% (£[Value])  

**P1 Metrics**:  
Weighted Fund Charge % = [Value]% (£[Value])  
Platform Charge % = [Value]% (£[Value])  
Ongoing Advice Fee % = [Value]% (£[Value])  
Discretionary Fund Manager Charge % = [Value]% (£[Value])  
Drawdown Fee % = [Value]% (£[Value])  
ProfitShare % = [Value]% (£[Value])  

**Total Annual Ongoing Charges**:  
- Royal London: [Total_%]% (£[Total_£])  
- P1: [Total_P1_%]% (£[Total_P1_£])  

**Comparison**: [Dynamic_Statement]  
---

**Rules**:  
- Use "0.0%" for missing values.  
- Format £ as £1,234.56.  
- No explanations. Only the template.
    """

    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
    )
    return response["choices"][0]["message"]["content"]
        
                        
                            

    
                       
'''def extract_iht_details_with_gpt(factfinding_text, product_report_text):
    """
    Use OpenAI GPT to extract Inheritance Tax (IHT) details by combining property values 
    from factfinding_text and pension information from product_report_text.
    Returns a string that is strictly JSON.
    """
    prompt = f"""
You are a financial assistant tasked with generating a comprehensive Inheritance Tax (IHT) calculation for a client in the UK.

---

### **Objective**:
Analyze the provided data and calculate the IHT liabilities, taking into account the following UK IHT rules and thresholds. Use the client’s address to estimate property value, search the `factfinding_text` for details about the client’s wife, dependents, mortgage, and debts, and include these details in the calculations.

---

### **IHT Rules and Calculations**:
1. **Inheritance Tax (IHT) Overview**:
   - Inheritance Tax is charged at 40% on the value of the taxable estate above the thresholds.

2. **Nil Rate Band (NRB)**:
   - Every individual has a Nil Rate Band (NRB) allowance of £325,000. This amount is **not taxed**.
   - For married couples or civil partners, the unused portion of the NRB can be transferred to the surviving spouse, effectively doubling the NRB to £650,000.

3. **Residence Nil Rate Band (RNRB)**:
   - If the estate includes a main residence passed to direct descendants (e.g., children or grandchildren), an additional **Residence Nil Rate Band (RNRB)** of up to £175,000 may apply.
   - The RNRB is transferable between spouses, effectively doubling to £350,000 for married couples or civil partners.
   - If there is **no property**, the RNRB is not applicable and should be set to £0.00.

4. **Property Ownership**:
   - Use the **client's address** from the FactFinding document to estimate their property value using external services such as property price estimators or by making an assumption based on standard UK property values for the area.
   - If no property is found or if the client does not own a property, the Residence Nil Rate Band (RNRB) is set to £0.00.

5. **Mortgage and Debts**:
   - Search the `factfinding_text` for any mortgage, loans, or debts.
   - Deduct these amounts from the total taxable estate to calculate the remaining estate.

6. **Wife and Dependents**:
   - Search the `factfinding_text` for information about the client's wife or civil partner and dependents.
   - If a wife or civil partner is found, include their NRB (£325,000) and RNRB (£175,000) in the calculations.
   - Use details about dependents (e.g., children or grandchildren) to determine eligibility for the RNRB.

7. **Pensions**:
   - Include pension values where applicable (e.g., drawdown pensions, lump sums) from the product report.

8. **Taxable Estate**:
   - Calculate the taxable estate as the sum of all assets (property, pensions, savings, investments) minus any applicable Nil Rate Bands and debts.

9. **Remaining Estate**:
   - If the taxable estate exceeds the combined Nil Rate Bands (NRB and RNRB), calculate the remaining estate subject to IHT.

10. **IHT Liability**:
   - Tax the remaining estate at 40% to determine the total IHT liability.

11. **Final Presentation**:
   - Ensure the table includes the following rows:
     - Property values (e.g., Main Residence)
     - Pension value
     - Total Taxable Estate
     - Mortgage and Debts
     - Nil Rate Band x2
     - Residence Nil Rate Band x2
     - Remaining Estate
     - Tax @ 40%
   - Use the format:
     - Amounts should be formatted as `£<value>.00`.
     - Ensure all calculations are rounded to two decimal places.

---

### **Input Data**:
1. FactFinding Document:
   {factfinding_text}

2. Royal London Report:
   {product_report_text}

---

### **Output Format**:
Return the IHT calculation in this JSON format:
```json
{
    "Property Details": [
        {"Description": "Main Residence", "Value": "£283,000.00"},
        {"Description": "Investment Property", "Value": "£250,000.00"}
    ],
    "Pension Details": {
        "Plan Value": "£124,029.82",
        "Review Period End Date": "31st December 2024"
    },
    "Dependents": [
        {"Name": "Child 1", "Relation": "Daughter"},
        {"Name": "Child 2", "Relation": "Son"}
    ],
    "Spouse": {
        "Name": "Jane Doe",
        "Eligible for RNRB": true
    },
    "Total Taxable Estate": "£657,029.82",
    "Mortgage and Debts": "£150,000.00",
    "Nil Rate Band Scenarios": {
        "Nil Rate Band x2": "£650,000.00",
        "Residence Nil Rate Band x2": "£350,000.00",
        "Remaining Estate": "£0.00",
        "Tax @ 40%": "£0.00"
    }
}

    """

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",  # or 'gpt-4' if you have access
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )

        raw_content = response["choices"][0]["message"]["content"]
        return raw_content

    except Exception as e:
        st.error("Error extracting IHT details: " + repr(e))'''

import openai
import csv

def extract_iht_details_with_gpt(factfinding_text, product_report_text):
    """
    Use OpenAI GPT to extract Inheritance Tax (IHT) details, but return CSV instead of JSON.
    The CSV will contain rows for each piece of data we want in the final table.
    """

    prompt = f"""
You are a financial assistant tasked with generating a comprehensive Inheritance Tax (IHT) calculation for a client in the UK.

---

### **Objective**:
Analyze the provided data and calculate the IHT liabilities, taking into account the following UK IHT rules and thresholds. Use the client’s address to estimate property value based on typical property prices for the area, search the `factfinding_text` for details about the client’s wife, dependents, mortgage, and debts, and include these details in the calculations.

---

### **IHT Rules and Calculations**:
1. **Inheritance Tax (IHT) Overview**:
   - Inheritance Tax is charged at 40% on the value of the taxable estate above the thresholds.

2. **Nil Rate Band (NRB)**:
   - Every individual has a Nil Rate Band (NRB) allowance of £325,000. This amount is **not taxed**.
   - For married couples or civil partners, the unused portion of the NRB can be transferred to the surviving spouse, effectively doubling the NRB to £650,000.

3. **Residence Nil Rate Band (RNRB)**:
   - If the estate includes a main residence passed to direct descendants (e.g., children or grandchildren), an additional **Residence Nil Rate Band (RNRB)** of up to £175,000 may apply.
   - The RNRB is transferable between spouses, effectively doubling to £350,000 for married couples or civil partners.
   - If there is **no property**, the RNRB is not applicable and should be set to £0.00.


4. **Property Ownership**:
   - Use the **client's address + zip code /postcode** from the FactFinding document to estimate their property value based on typical UK property prices for the area.
   - The address provided is:
     ```
     {"Address + zipcode/postcode"}
     ```
   - Based on this address, estimate the current market value of the property using knowledge of UK property markets up to October 2023.
   - **Example**:
     - *For instance, if the address is:* example address
       ```
       24 Wesley Road
       Leonard Stanley
       Stonehouse
       Gloucestershire
       GL10 3PF
       ```
       *you would consider the typical property prices in the GL10 postcode area, factoring in property type, location desirability, and market trends to estimate a realistic value.*

   - If no property is found or if the client does not own a property, set the Residence Nil Rate Band (RNRB) to £0.00.

5. **Mortgage and Debts**:
   - Search the `factfinding_text` for any mortgage, loans, or debts.
   - Deduct these amounts from the total taxable estate to calculate the remaining estate.

6. **Wife and Dependents**:
   - Search the `factfinding_text` for information about the client's wife or civil partner and dependents.
   - If a wife or civil partner is found, include their NRB (£325,000) and RNRB (£175,000) in the calculations.
   - Use details about dependents (e.g., children or grandchildren) to determine eligibility for the RNRB.

7. **Pensions**:
   - Include pension values where applicable (e.g., drawdown pensions, lump sums) from the product report.

8. **Taxable Estate**:
   - Calculate the taxable estate as the sum of all assets (property, pensions, savings, investments) minus any applicable Nil Rate Bands and debts.

9. **Remaining Estate**:
   - If the taxable estate exceeds the combined Nil Rate Bands (NRB and RNRB), calculate the remaining estate subject to IHT.

10. **IHT Liability**:
   - Tax the remaining estate at 40% to determine the total IHT liability.

11. **Final Presentation**:
   - Ensure the table includes the following rows:
     - Property values (e.g., Main Residence)
     - Pension value
     - Total Taxable Estate
     - Mortgage and Debts
     - Nil Rate Band x2
     - Residence Nil Rate Band x2
     - Remaining Estate
     - Tax @ 40%
   - Use the format:
     - Amounts should be formatted as `£<value>.00`.
     - Ensure all calculations are rounded to two decimal places.
   - Dont provide explaination, just provide the results.
---

### **Input Data**:
1. FactFinding Document:
   {factfinding_text}

2. Royal London Report:
   {product_report_text}

### **Output Format**
Return only bullet points with short sentences, formatted as follows:

- Main Residence worth = £253,000.00
- Pension worth = £124,029.82
- Total Taxable Estate = £407,029.82
- Mortgage and Debts = £0.00
- Nil Rate Band x2 = £650,000.00
- Residence Nil Rate Band x2 = £350,000.00
- Remaining Estate = £0.00
- Tax @ 40% = £0.00

    """

    try:
        
        
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        bullet_points = response["choices"][0]["message"]["content"].strip()
        return bullet_points  # Return as text string
    except Exception as e:
        raise RuntimeError("Error extracting IHT details as bullet points: " + repr(e))

from io import StringIO


    
def create_plan_report_table(document, table_data):
    if not isinstance(table_data, list) or not all(isinstance(plan, dict) for plan in table_data):
        raise ValueError("Invalid table_data format")

    document.add_heading("Plan Report Details", level=2)
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    headers = ["Provider", "Plan Number", "Plan Type", "Current Value", "Recommendation"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header

    total_value = 0
    for plan in table_data:
        row = table.add_row().cells
        row[0].text = plan.get("Provider", "")
        row[1].text = plan.get("Plan Number", "")
        row[2].text = plan.get("Plan Type", "")
        row[3].text = plan.get("Current Value", "")
        row[4].text = ""

        try:
            current_value = float(plan["Current Value"].replace("£", "").replace(",", ""))
            total_value += current_value
        except (KeyError, ValueError):
            pass

    total_row = table.add_row().cells
    total_row[2].text = "Total"
    total_row[3].text = f"£{total_value:,.2f}"
    return table


def create_comparison_table(document, sap_comparison_table):
    """
    Generate the 'Comparison at Age {age}' table in the Word document.
    """
    if not sap_comparison_table or "Table" not in sap_comparison_table:
        raise ValueError("SAP comparison table data is missing or incorrectly formatted.")

    table_data = sap_comparison_table["Table"]

    # Validate header and rows
    if "Assumed Growth Rates" not in table_data or len(table_data["Assumed Growth Rates"]) == 0:
        raise ValueError("Table header is missing or incomplete.")

    # Determine number of rows and columns
    header = [""] + table_data["Assumed Growth Rates"]  # Add an empty cell for the first column
    rows = [
        ["Existing Schemes"] + table_data["Existing Schemes"],
        ["P1 Pension Account"] + table_data["P1 Pension Account"],
        ["Rate of Return Required from P1"] + table_data["Rate of Return Required from P1"],
        ["Effect on Fund if Moved to P1"] + table_data["Effect on Fund if Moved to P1"],
        ["Reduction in Yield if Moved to P1"] + table_data["Reduction in Yield if Moved to P1"],
    ]

    # Add the table to the document
    table = document.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = "Table Grid"

    # Fill in the header row
    for col_idx, header_text in enumerate(header):
        table.cell(0, col_idx).text = header_text

    # Fill in the data rows
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = value

'''def create_annuity_quotes_table(document, annuity_quotes):
    """
    Generate a table with annuity quotes and append it to the provided document.
    """
    if not annuity_quotes or "Quotes" not in annuity_quotes:
        raise ValueError("Invalid annuity quotes data.")

    quotes = annuity_quotes["Quotes"]

    # Add a heading for the table
    document.add_heading("Annuity Quotes", level=2)

    # Create the table with rows for attributes and a header for quotes
    table = document.add_table(rows=5, cols=len(quotes) + 1)
    table.style = 'Table Grid'

    # Fill the first cell in the header with an empty label
    table.cell(0, 0).text = ""

    # Add headers for Quotes (Quote 1, Quote 2, etc.)
    for idx in range(len(quotes)):
        table.cell(0, idx + 1).text = f"Quote {idx + 1}"

    # Define the rows for attributes
    attributes = ["Purchase Amount", "Monthly Amount", "Yearly Amount", "Yearly Increase"]

    for row_idx, attribute in enumerate(attributes, start=1):
        # Add the attribute name to the first column
        table.cell(row_idx, 0).text = attribute
        # Fill in the data for each quote
        for col_idx, quote in enumerate(quotes):
            table.cell(row_idx, col_idx + 1).text = quote.get(attribute, "")

    return table'''


 

def create_new_document(template_path, factfinding_text, risk_details, table_data, product_report_text, plan_report_text,fund_performance_data,dark_star_performance_data,sap_comparison_table,annuity_quotes_text,fund_comparison_text,iht_bullet_points, output_path):
    """
    Create a well-formatted document by replacing placeholders, appending tables,
    and inserting dynamically generated sections while preserving static text.
    """


 # Parse client details from GPT response
    client_details_json_str = extract_client_details_with_gpt(factfinding_text)
    try:
        client_details = extract_client_details_with_gpt(factfinding_text)
    except json.JSONDecodeError as e:
        raise ValueError(
            "Error parsing client details JSON: " + repr(e) +
            "\nRaw GPT Response:\n" + client_details_json_str
        )
    # Validate risk details
    if not isinstance(risk_details, dict):
        raise ValueError(f"Expected risk_details to be a dictionary, but got: {type(risk_details)}")

    # Generate dynamic sections
    current_situation = generate_current_situation(factfinding_text)
    priorities_and_objectives = generate_priorities_and_objectives(factfinding_text)
    attitude_to_risk = f"""
    {client_details.get('Full name', 'Client')}, the risk profiler completed with you came out with an attitude to risk level {risk_details.get('risk_level', 'N/A')} which is {risk_details.get('risk_type', 'N/A')}. 
    {risk_details.get('first_sentence', '')} {risk_details.get('last_sentence', '')}
    """
    pension_review_section = generate_pension_review_section(product_report_text)
    swr_section = generate_safe_withdrawal_rate_section(plan_report_text)
    ''''
    age = sap_comparison_table.get("Age", "N/A")  # Default to "N/A" if Age is not found
# Extract middle value for text generation
    middle_value_str = sap_comparison_table["Table"]["Rate of Return Required from P1"][1]  # Second column
    try:
        middle_value = float(middle_value_str.replace('%', '').strip())
    except ValueError:
        raise ValueError(f"Invalid middle value extracted: {middle_value_str}")
    
    

    # Determine the text based on the value
    if middle_value < 0.0:
        below_table_text = f"The critical yield required to match the benefits of your current scheme is {middle_value_str}, indicating that the proposed arrangement would need less performance of {middle_value_str} per annum to make up the costs of transferring. This is because the proposed arrangement is more cost-effective than your current arrangement."
    elif 0.0 <= middle_value < 3.0:
        below_table_text = f"The critical yield required to match the benefits of your current scheme is {middle_value_str}, indicating that the proposed arrangement would need an additional fund performance of {middle_value_str} per annum to make up the costs of transferring.\n\nI believe the chosen fund will be able to achieve this over the long term, although this is not guaranteed."
    else:  # middle_value >= 3.0
        below_table_text = f"The critical yield required to match the benefits of your current scheme is {middle_value_str}, indicating that the proposed arrangement would need an additional performance of {middle_value_str} per annum to make up the costs of transferring.\n\nI can not guarantee that the recommended fund can match the additional performance required to make up the costs of transferring, but I still believe that transferring out is in your best interests. Performance is only one consideration to make when transferring out."
'''
    # Load the template
    original_doc = Document(template_path)
    new_doc = Document()

  
    # Replace placeholders while preserving static text
    for paragraph in original_doc.paragraphs:
        text = paragraph.text

        # Replace placeholders dynamically
        if "{Full name}" in text:
            text = text.replace("{Full name}", client_details.get("Full name", ""))
        if "{Address}" in text:
            text = text.replace("{Address}", client_details.get("Address", ""))
        if "{Today’s date}" in text:
            text = text.replace("{Today’s date}", client_details.get("Today’s date", ""))
        if "{salutation}" in text:
            text = text.replace("{salutation}", client_details.get("salutation", ""))
        if "{Current_Situation}" in text:
            text = text.replace("{Current_Situation}", current_situation)
        if "{Priorities_and_Objectives}" in text:
            text = text.replace("{Priorities_and_Objectives}", priorities_and_objectives)
        if "{Attitude_to_Risk}" in text:
            text = text.replace("{Attitude_to_Risk}", attitude_to_risk)
        if "{Review of Existing Royal London Personal Pension}" in text:
            text = text.replace("{Review of Existing Royal London Personal Pension}", pension_review_section)
        if "{Safe Withdrawal Rate (SWR)}" in text:
            text = text.replace("{Safe Withdrawal Rate (SWR)}", swr_section)

        # Handle the {table} placeholder
        if "{table1}" in text:
            paragraph.text = "Overview of Current Investments\nHere is a breakdown of your current investments."
            create_plan_report_table(new_doc, table_data)  # Insert table right after the paragraph
            continue
        # Handle the {table2-1} placeholder for fund performance
        if "{table2-1}" in text:
            bullet_points = "Extracted Fund Performance\n\n"
            for fund in fund_performance_data:
                bullet_points += f"**{fund['Fund']}**\n"
                for year in range(1, 6):
                    year_key = f"Year {year}"
                    benchmark_key = fund.get("Benchmark", {}).get(year_key, "N/A")
                    year_value = fund.get(year_key, "N/A")
                    bullet_points += f"- {year_key}: {year_value} (Benchmark: {benchmark_key})\n"
                cumulative_performance = fund.get("Cumulative (5 YR)", "N/A")
                cumulative_benchmark = fund.get("Benchmark", {}).get("Cumulative (5 YR)", "N/A")
                bullet_points += f"- Cumulative 5-Year Performance: {cumulative_performance} (Benchmark: {cumulative_benchmark})\n\n"
            text = text.replace("{table2-1}", bullet_points.strip())

                # Handle the {table2-2} placeholder for Dark Star performance
        if "{table2-2}" in text:
            bullet_points = "Extracted Dark Star Performance\n\n"
            for fund in dark_star_performance_data:
                bullet_points += f"**{fund['Fund']}**\n"
                for year in range(1, 6):
                    year_key = f"Year {year}"
                    year_value = fund.get(year_key, "N/A")
                    bullet_points += f"- {year_key}: {year_value}\n"
                cumulative_performance = fund.get("Cumulative (5 YR)", "N/A")
                bullet_points += f"- Cumulative 5-Year Performance: {cumulative_performance}\n\n"
            text = text.replace("{table2-2}", bullet_points.strip())   

                # Handle placeholders
        # Handle the {table3-1} placeholder for SAP comparison
        if "{table3-1}" in text:
            heading = f"Comparison at Age {sap_comparison_table.get('Age', 'N/A')}"
            paragraph_before_table = (
                f"The table below shows the projected value of your pensions at the age of "
                f"{sap_comparison_table.get('Age', 'N/A')}, firstly if it were to remain in your "
                f"current arrangement and secondly were it to be transferred to P1."
            )
            new_doc.add_heading(heading, level=2)
            new_doc.add_paragraph(paragraph_before_table)
            create_comparison_table(new_doc, sap_comparison_table)
            new_doc.add_paragraph("")  # Add an empty paragraph to create a blank line
            ''' new_doc.add_paragraph(below_table_text)'''  # Add the generated text below the table

            continue

        if "{Annuity_Quotes}" in text:
                    if annuity_quotes_text:
                        text = text.replace("{Annuity_Quotes}", annuity_quotes_text)
                    else:
                        text = text.replace("{Annuity_Quotes}", "No annuity quotes available.") 

                # Replace placeholder for Fund Comparison
        if "{Fund_Comparison}" in text:
            text = text.replace("{Fund_Comparison}", fund_comparison_text)  

# Handle the {IHT_Table} placeholder with bullet points
        if "{IHT_Table}" in text:
            new_doc.add_heading("Inheritance Tax (IHT) Details", level=2)
            if not iht_bullet_points or not isinstance(iht_bullet_points, str):
                # If bullet points are missing or invalid
                new_doc.add_paragraph("No valid IHT details found.")
            else:
                # Insert bullet points as a bulleted list
                for line in iht_bullet_points.split('\n'):
                    if line.strip().startswith("-"):
                        new_doc.add_paragraph(line.strip(), style='List Bullet')
                    else:
                        new_doc.add_paragraph(line.strip())
        # Add formatted paragraph to the new document
        new_paragraph = new_doc.add_paragraph(text)
        new_paragraph.style = new_doc.styles['Normal']
        new_paragraph.paragraph_format.space_before = Pt(6)
        new_paragraph.paragraph_format.space_after = Pt(6)
        new_paragraph.paragraph_format.line_spacing = Pt(12)  # Set consistent line spacing
        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Adjust document margins
    section = new_doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Save the combined document
    new_doc.save(output_path)
    print(f"Debug: Document saved successfully at {output_path}")


    return swr_section


def save_uploaded_file(uploaded_file, folder):
    """
    Save the uploaded file to a specified folder.
    Args:
    - uploaded_file: File uploaded through UI.
    - folder (str): Folder to save the file.
    Returns:
    - str: File path of the saved file.
    """
    os.makedirs(folder, exist_ok=True)
    file_path = os.path.join(folder, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.read())
    return file_path
